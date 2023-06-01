import io
import os
import re
import sys
from configparser import ConfigParser

from PySide6 import QtCore
from PySide6.QtCore import QEventLoop, QTimer
from PySide6.QtWidgets import QApplication, QMainWindow, QFileDialog
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from openpyxl.reader.excel import load_workbook
from win32com.client import DispatchEx

from docrobot.form import Ui_MainWindow


class Project(object):
    p_comname = ''  # 公司名称
    p_order = ''  # 序号
    p_name = ''  # 项目名
    p_start = ''
    p_end = ''
    p_cost = ''
    p_people = ''  # 人数
    p_owner = ''  # 项目负责人
    p_rnd = ''  # 研发人员
    p_money = ''  # 总预算


class EmittingStr(QtCore.QObject):
    textWritten = QtCore.Signal(str)

    def write(self, text):
        self.textWritten.emit(str(text))
        loop = QEventLoop()
        QTimer.singleShot(100, loop.quit)
        loop.exec()
        QApplication.processEvents()


class MainWindow(Ui_MainWindow, QMainWindow):
    workdir = ''
    file_prj = ''
    file_pat = ''
    pat_dict = {}  # 专利->序号字典
    pat_dict2 = {}  # 专利->专利编号字典
    arr_prj = []

    def __init__(self):
        super(MainWindow, self).__init__()
        self.setupUi(self)
        sys.stdout = EmittingStr()
        sys.stdout.textWritten.connect(self.outputWritten)
        sys.stderr = EmittingStr()
        sys.stderr.textWritten.connect(self.outputWritten)

        self.actionSelect_Dir.triggered.connect(self.setDocUrl)
        self.actioncheck.triggered.connect(lambda: self.checkpatent(True))
        self.actionreplace.triggered.connect(lambda: self.replaceprj(True))
        self.actioncheckall.triggered.connect(self.checkall)

        self.config = ConfigParser()
        try:
            self.config.read('config.ini', encoding='UTF-8')
            self.workdir = self.config['config']['lasting']
        except KeyError:
            self.config.add_section('config')
        try:
            if self.workdir != '':
                self.onchangeworkdir()
        except FileNotFoundError as e:
            self.textEdit.append('路径错误：' + e.filename)
            self.workdir = ''
        self.lineEdit.setText(self.workdir)

    def outputWritten(self, text):
        self.textEdit.append(text.strip())

    def setDocUrl(self):
        # 重新选择输入和输出目录时，进度条设置为0，文本框的内容置空
        tempdir = QFileDialog.getExistingDirectory(self, "选中项目所在目录", r"")
        if tempdir != '':
            self.workdir = tempdir
            self.onchangeworkdir()
            with open('config.ini', 'w', encoding='utf-8') as file:
                self.config['config']['lasting'] = self.workdir
                self.config.write(file)  # 数据写入配置文件

    def onchangeworkdir(self):
        self.lineEdit.setText(self.workdir)
        for file_sum in os.listdir(self.workdir):
            if file_sum.endswith('立项报告汇总表.xlsx') and not file_sum.startswith('~$'):
                self.file_prj = self.workdir + '/' + file_sum
            if file_sum.endswith('知识产权汇总表.xlsx') and not file_sum.startswith('~$'):
                self.file_pat = self.workdir + '/' + file_sum
        if self.file_prj == '':
            self.textEdit.append('没找到：' + '立项报告汇总表.xlsx')
        if self.file_pat == '':
            self.textEdit.append('没找到：' + '知识产权汇总表.xlsx')

    def replaceprj(self, modify=False):
        if modify:
            self.textEdit.setText('')
            self.update_data()

        for project in self.arr_prj:
            # self.textEdit.append(project.p_order + ' 项目开始处理...')
            match = 0  # 已匹配条目数
            doc_name = ''
            try:
                doc_name = self.workdir + '/RD' + project.p_order + project.p_name + '.docx'
                document = Document(doc_name)
                # debug_doc(document)
                # TODO to be fixed
                # replace_header(document)
                match = match + self.first_table(document, project)
                match = match + self.start_time(document, project)
                match = match + self.second_table(document, project)
                match = match + self.third_table(document, project)

                match = match + self.checkpat2(document, project)

                if modify:
                    document.save(doc_name)
            except PackageNotFoundError:
                self.textEdit.append('Error打开文件错误：' + doc_name)
            except PermissionError:
                self.textEdit.append('Error 保存文件错误，可能是文件已被打开：' + doc_name)
            self.textEdit.append(project.p_order + ' 项目：处理完成。 <font color="green"><b>'
                                 + str(match) + ' </b></font> 项条目匹配。')

    def checkpatent(self, modify=False):
        if modify:
            self.textEdit.setText('')
            self.update_data()

        match = 0  # 已匹配条目数
        changed = False
        wb = load_workbook(self.file_prj)
        ws = wb.active
        max_row_num = ws.max_row
        range_cell = ws[f'A3:P{max_row_num}']
        i: int = 0
        for r in range_cell:
            if r[11].value is None:
                break
            pat_name = str(r[11].value).strip()

            if pat_name == '无':
                if r[14].value != '无':
                    self.textEdit.append(str(i + 3) + ' 行: ' + str(r[14].value) + ' ===> ' + '无')
                    r[14].value = pat_name
                    changed |= True
                else:
                    match = match + 1
            else:
                lst = pat_name.splitlines()
                rep = [self.pat_dict[x] if x in self.pat_dict else x for x in lst]
                for element in rep:
                    if re.search('^\d\d$', element) is None:
                        self.textEdit.append('专利IP错误：' + element)
                rep = map(lambda e: 'IP' + e, rep)
                new_ip = ';'.join(rep)
                if r[14].value != new_ip:
                    self.textEdit.append(str(i + 3) + ' 行: ' + str(r[14].value) + ' ===> ' + new_ip)
                    r[14].value = new_ip
                    changed |= True
                else:
                    match = match + 1
            i = i + 1
        try:
            if changed and modify:
                wb.save(self.file_prj)
                xl_app = DispatchEx("Excel.Application")
                xl_app.Visible = False
                xl_app.DisplayAlerts = False
                xl_book = xl_app.Workbooks.Open(self.file_prj)
                xl_book.Saved = False
                xl_book.Close(True)
                xl_book = None
                xl_app.Quit()
                xl_app = None
        except PermissionError:
            self.textEdit.append('写文件失败，关闭其他占用该文件的程序.' + self.file_prj)
        wb.close()
        self.textEdit.append('项目立项表IP更新完成。 <font color="green"><b>' + str(match) + ' </b></font> 项条目匹配。')

    def checkpat2(self, doc, prj):
        match = 0
        lst = prj.pat_list.splitlines()
        for pat_name in lst:
            if pat_name in self.pat_dict2:
                pat_num = self.pat_dict2[pat_name]
                found = False
                regex = pat_name.replace('[', r'\[').replace(']', r'\]'.replace('(', r'\(').replace(')', r'\)'))
                for i, para in enumerate(doc.tables[2].rows[4].cells[0].paragraphs):
                    if re.search(regex + '，.*号：', para.text):
                        found |= True
                        result = re.search(regex + '，.*号：' + pat_num, para.text)
                        if result is None:
                            self.textEdit.append('专利名和编号不匹配：' + pat_name + ' , ' + pat_num)
                            self.textEdit.append('文档内容：' + para.text)
                        else:
                            match = match + 1
                if not found:
                    self.textEdit.append('全文找不到：' + pat_name)
            else:
                self.textEdit.append('Error没有找到专利：' + pat_name)
        return match

    def update_data(self):
        self.pat_dict.clear()
        self.pat_dict2.clear()
        with open(self.file_pat, "rb") as f:
            in_mem_file = io.BytesIO(f.read())
        wb = load_workbook(in_mem_file, read_only=True, data_only=True)
        ws = wb.active
        if ws is not None:
            max_row_num = ws.max_row
            range_cell = ws[f'A3:D{max_row_num}']
            for r in range_cell:
                if r[0].value is None:
                    break
                p_order = str(r[0].value).strip().zfill(2)
                p_name = str(r[1].value).strip()
                p_patnum = str(r[3].value).strip()
                self.pat_dict[p_name] = p_order
                self.pat_dict2[p_name] = p_patnum
        else:
            self.textEdit.append('Error ' + self.file_pat + ' 文件格式错误。')
        wb.close()

        self.arr_prj.clear()
        com_name = 'XX公司'
        with open(self.file_prj, "rb") as f:
            in_mem_file = io.BytesIO(f.read())
        wb = load_workbook(in_mem_file, read_only=True, data_only=True)
        ws = wb.active
        if ws is not None:
            if str(ws['A1'].value).find(u'公司') != -1:
                com_name = str(ws['A1'].value).split("公司")[0] + '公司'
            else:
                self.textEdit.append("Error: 找不到 公司名")
            max_row_num = ws.max_row
            range_cell = ws[f'A3:P{max_row_num}']
            for r in range_cell:
                if r[0].value is None:
                    break
                project = Project()
                project.p_comname = com_name
                project.p_order = str(r[0].value).strip().zfill(2)
                project.p_name = str(r[1].value).strip()  # 项目名称
                project.p_start = r[2].value.strftime('%Y-%m-%d')
                project.p_end = r[3].value.strftime('%Y-%m-%d')
                project.p_cost = str(r[5].value).strip()
                project.p_people = str(r[6].value).strip()  # 人数
                project.p_owner = str(r[7].value).strip()  # 项目负责人
                project.p_rnd = str(r[8].value).strip()  # 研发人员
                project.p_money = str(r[9].value).strip()  # 总预算
                project.pat_list = str(r[11].value).strip()  # 知识产权名称
                project.ip_list = str(r[14].value).strip()  # IP
                self.arr_prj.append(project)
        else:
            self.textEdit.append('Error ' + self.file_prj + ' 文件格式错误。')
        wb.close()
        if len(self.arr_prj) == 0:
            self.textEdit.append("Error: 立项汇总表错误，使用Excel重新保存.")

    def checkall(self):
        self.textEdit.setText('')
        self.update_data()

        self.checkpatent(False)
        self.replaceprj(False)

    def check_and_change(self, doc, replace):
        """
        遍历word中的所有 paragraphs，在每一段中发现含有key 的内容，就替换为 value 。
        （key 和 value 都是replace_dict中的键值对。）
        """
        for para in doc.paragraphs:
            for i in range(len(para.runs)):
                for key, value in replace.items():
                    if key in para.runs[i].text:
                        self.textEdit.append(key + "-->" + value)
                        para.runs[i].text = para.runs[i].text.replace(key, value)
        return doc

    def replace_tables(self, doc, replace):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for i in range(len(para.runs)):
                            # self.textEdit.append(">>>" + para.runs[i].text)
                            for key, value in replace.items():
                                if key in para.runs[i].text:
                                    self.textEdit.append(key + "-->" + value)
                                    para.runs[i].text = para.runs[i].text.replace(key, value)
        return doc

    @staticmethod
    def clear_runs(runs):
        for i, run in enumerate(runs):
            if i > 0:
                run.clear()
        return runs

    def debug_doc(self, doc):
        for i, sect in enumerate(doc.sections):
            for j, para in enumerate(sect.header.paragraphs):
                self.textEdit.append(f'Sec.{i} Para.{j} : ', para.text, sep='')
                for k, run in enumerate(para.runs):
                    self.textEdit.append(f'Sec.{i} Para.{j} Run{k}: ', run.text, sep='')
        for i, para in enumerate(doc.paragraphs):
            self.textEdit.append(f'Para.{i} : ', para.text, sep='')
            for j, run in enumerate(para.runs):
                self.textEdit.append(f'Para.{i} Run{j}: ', run.text, sep='')
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    for l, para in enumerate(cell.paragraphs):
                        self.textEdit.append(f'Table.{i} Row.{j} Cell{k} Para.{l} : ', para.text, sep='')
                        # for j, run in enumerate(para.runs):
                        #     self.textEdit.append(f'Para.{i} Run{j}: ', run.text, sep='')

    def replace_header(self, doc, prj):
        self.check_replace(doc.sections[0].header.paragraphs, '.*公司', prj.com_name)

    def first_table(self, doc, prj):
        match = 0
        if doc.tables[0].rows[0].cells[0].paragraphs[0].text == '项目名称：':
            doc.tables[0].rows[0].cells[1].paragraphs[0].runs[0].text = prj.p_name
            self.clear_runs(doc.tables[0].rows[0].cells[1].paragraphs[0].runs)
            match = match + 1
        if doc.tables[0].rows[1].cells[0].paragraphs[0].text == '项目编号：':
            doc.tables[0].rows[1].cells[1].paragraphs[0].runs[0].text = prj.p_start[0:4] + 'RD' + prj.p_order
            self.clear_runs(doc.tables[0].rows[1].cells[1].paragraphs[0].runs)
            match = match + 1
        if doc.tables[0].rows[2].cells[0].paragraphs[0].text == '项目负责人：' and prj.p_owner != 'None':
            doc.tables[0].rows[2].cells[1].paragraphs[0].runs[0].text = prj.p_owner
            self.clear_runs(doc.tables[0].rows[2].cells[1].paragraphs[0].runs)
            match = match + 1
        if doc.tables[0].rows[3].cells[0].paragraphs[0].text == '项目周期：':
            doc.tables[0].rows[3].cells[1].paragraphs[0].runs[0].text = prj.p_start + '至' + prj.p_end
            self.clear_runs(doc.tables[0].rows[3].cells[1].paragraphs[0].runs)
            match = match + 1
        return match

    def start_time(self, doc, prj):
        match = self.check_replace(doc.paragraphs, '申请立项时间：\d{4}[-/]\d{1,2}[-/]\d{1,2}',
                                   '申请立项时间：' + prj.p_start)
        return match

    def second_table(self, doc, prj):
        match = 0
        if doc.tables[1].rows[0].cells[0].paragraphs[0].text == '项目立项名称':
            doc.tables[1].rows[0].cells[1].paragraphs[0].runs[0].text = prj.p_name
            self.clear_runs(doc.tables[1].rows[0].cells[1].paragraphs[0].runs)
            match = match + 1

        match = match + self.check_replace(doc.tables[1].rows[1].cells[1].paragraphs
                                           , '项目团队由(.*)人组成，项目实施周期为(.*)个月。'
                                           ,
                                           '项目团队由' + prj.p_people + '人组成，项目实施周期为' + prj.p_cost + '个月。')
        match = match + self.check_replace(doc.tables[1].rows[6].cells[1].paragraphs
                                           , '\d{4}[-/]\d{1,2}[-/]\d{1,2}至\d{4}[-/]\d{1,2}[-/]\d{1,2}',
                                           prj.p_start + '至' + prj.p_end)
        match = match + self.check_replace(doc.tables[1].rows[7].cells[1].paragraphs
                                           , '项目总资金预算.*万元', '项目总资金预算' + prj.p_money + '万元')

        match = match + self.check_replace(doc.tables[1].rows[8].cells[1].paragraphs
                                           , '项目总人数：.*人', '项目总人数：' + prj.p_people + '人')
        if prj.p_owner != 'None':
            match = match + self.check_replace(doc.tables[1].rows[8].cells[1].paragraphs, '项目负责人：.*',
                                               '项目负责人：' + prj.p_owner)
        if prj.p_rnd != 'None':
            match = match + self.check_replace(doc.tables[1].rows[8].cells[1].paragraphs, '研发成员：.*',
                                               '研发成员：' + prj.p_rnd)
        match = match + self.check_replace(doc.tables[1].rows[9].cells[1].paragraphs, '\d{4}[-/]\d{1,2}[-/]\d{1,2}',
                                           prj.p_start)

        return match

    def check_replace(self, paras, regex, dst):
        match = 0
        for i, para in enumerate(paras):
            result = re.search(regex, para.text)
            if result is not None:
                if result.group() != dst:
                    self.textEdit.append(result.group() + ' ===> ' + dst)
                    para.runs[0].text = re.sub(regex, dst,
                                               para.text)
                    self.clear_runs(para.runs)
                match = match + 1
                break  # 只替换一次就够用
        return match

    def third_table(self, doc, prj):
        match = 0
        if doc.tables[2].rows[0].cells[0].paragraphs[0].text == '项目名称':
            doc.tables[2].rows[0].cells[1].paragraphs[0].runs[0].text = prj.p_name
            self.clear_runs(doc.tables[1].rows[0].cells[1].paragraphs[0].runs)
            match = match + 1
        match = match + self.check_replace(doc.tables[2].rows[1].cells[1].paragraphs
                                           , '\d{4}[-/]\d{1,2}[-/]\d{1,2}', prj.p_end)
        match = match + self.check_replace(doc.tables[2].rows[2].cells[1].paragraphs
                                           , '\d{4}[-/]\d{1,2}[-/]\d{1,2}至\d{4}[-/]\d{1,2}[-/]\d{1,2}',
                                           prj.p_start + '至' + prj.p_end)

        if prj.p_owner != 'None':
            doc.tables[2].rows[3].cells[1].paragraphs[0].runs[0].text = prj.p_owner
            self.clear_runs(doc.tables[2].rows[3].cells[1].paragraphs[0].runs)
            match = match + 1
        return match


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
