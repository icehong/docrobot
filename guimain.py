import os
import re
import sys
from configparser import ConfigParser

from PySide6 import QtCore, QtGui
from PySide6.QtCore import QEventLoop, QTimer
from PySide6.QtWidgets import QApplication, QMainWindow, QFileDialog
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from openpyxl.reader.excel import load_workbook
from win32com.client import Dispatch

import main
from form import Ui_MainWindow


class EmittingStr(QtCore.QObject):
    textWritten = QtCore.Signal(str)

    def write(self, text):
        self.textWritten.emit(str(text))
        loop = QEventLoop()
        QTimer.singleShot(100, loop.quit)
        loop.exec()
        QApplication.processEvents()


class MainWindow(Ui_MainWindow, QMainWindow):
    workdir = ''
    file_prj = ''
    file_pat = ''
    pat_dict = {}  # 专利->序号字典
    pat_dict2 = {}  # 专利->专利编号字典
    arr_prj = []

    def __init__(self):
        super(MainWindow, self).__init__()
        self.setupUi(self)
        sys.stdout = EmittingStr()
        sys.stdout.textWritten.connect(self.outputWritten)
        sys.stderr = EmittingStr()
        sys.stderr.textWritten.connect(self.outputWritten)

        self.actionSelect_Dir.triggered.connect(self.setDocUrl)
        self.actioncheck.triggered.connect(self.checkpatent)
        self.actionreplace.triggered.connect(self.replaceprj)

        self.config = ConfigParser()
        try:
            self.config.read('config.ini', encoding='UTF-8')
            self.workdir = self.config['config']['lasting']
            self.onchangeworkdir()
        except:
            self.config.add_section('config')
            pass
        self.lineEdit.setText(self.workdir)

    def outputWritten(self, text):
        cursor = self.textEdit.textCursor()
        cursor.movePosition(QtGui.QTextCursor.End)
        cursor.insertText(text)
        self.textEdit.setTextCursor(cursor)
        self.textEdit.ensureCursorVisible()

    def setDocUrl(self):
        # 重新选择输入和输出目录时，进度条设置为0，文本框的内容置空
        tempdir = QFileDialog.getExistingDirectory(self, "选中项目所在目录", r"")
        if tempdir != '':
            self.workdir = tempdir
            self.onchangeworkdir()
            with open('config.ini', 'w', encoding='utf-8') as file:
                self.config['config']['lasting'] = self.workdir
                self.config.write(file)  # 数据写入配置文件

    def onchangeworkdir(self):
        self.lineEdit.setText(self.workdir)
        for file_sum in os.listdir(self.workdir):
            if file_sum.endswith('立项报告汇总表.xlsx') and not file_sum.startswith('~$'):
                self.file_prj = self.workdir + '/' + file_sum
            if file_sum.endswith('知识产权汇总表.xlsx') and not file_sum.startswith('~$'):
                self.file_pat = self.workdir + '/' + file_sum
        if self.file_prj == '':
            self.textEdit.append('没找到：' + '立项报告汇总表.xlsx')
        if self.file_pat == '':
            self.textEdit.append('没找到：' + '知识产权汇总表.xlsx')

    def replaceprj(self):
        self.upate_global()
        wb = load_workbook(self.file_prj, read_only=True, data_only=True)
        ws = wb.active
        if str(ws['A1'].value).find(u'公司') != -1:
            com_name = str(ws['A1'].value).split("公司")[0] + '公司'
        else:
            print("Error: 找不到 公司名")
            return

        max_row_num = ws.max_row
        rangeCell = ws[f'A3:P{max_row_num}']
        for r in rangeCell:
            if r[0].value is None:
                break
            project = main.Project()
            project.p_comname = com_name
            project.p_order = str(r[0].value).strip().zfill(2)
            project.p_name = str(r[1].value).strip()
            project.p_start = r[2].value.strftime('%Y-%m-%d')
            project.p_end = r[3].value.strftime('%Y-%m-%d')
            project.p_cost = str(r[5].value).strip()
            project.p_people = str(r[6].value).strip()  # 人数
            project.p_owner = str(r[7].value).strip()  # 项目负责人
            project.p_rnd = str(r[8].value).strip()  # 研发人员
            project.p_money = str(r[9].value).strip()  # 总预算
            self.textEdit.append('开始处理项目：' + project.p_order)

            try:
                doc_name = self.workdir + '/RD' + project.p_order + project.p_name + '.docx'
                document = Document(doc_name)
                # debug_doc(document)
                # TODO to be fixed
                # replace_header(document)
                main.first_table(document, project)
                main.start_time(document, project)
                main.second_table(document, project)
                main.third_table(document, project, self.pat_dict2)
                document.save(doc_name)
            except PackageNotFoundError:
                self.textEdit.append('Error打开文件错误：' + doc_name)
        self.textEdit.append('检查和更新完成.')

    def checkpatent(self):
        self.upate_global()

        wb = load_workbook(self.file_prj, rich_text=True)
        ws = wb.active
        max_row_num = ws.max_row
        rangeCell = ws[f'A3:P{max_row_num}']
        i: int = 0
        for r in rangeCell:
            if r[11].value is None:
                break
            pat_name = str(r[11].value).strip()
            p_order = self.arr_prj[i].p_order
            p_name = self.arr_prj[i].p_name

            if pat_name == '无':
                if r[14].value != '无':
                    self.textEdit.append(str(i + 3) + ' 行: ' + str(r[14].value) + ' 替换为 ' + pat_name)
                    r[14].value = pat_name
            else:
                lst = pat_name.splitlines()
                for pat in lst:
                    if pat in self.pat_dict2:
                        self.checkpat2(self.workdir + '/RD' + p_order + p_name + '.docx', pat, self.pat_dict2[pat])
                    else:
                        self.textEdit.append('Error没有找到专利：' + pat)

                rep = [self.pat_dict[x] if x in self.pat_dict else x for x in lst]
                for element in rep:
                    if re.search('^\d\d$', element) is None:
                        self.textEdit.append('Error没有找到专利：' + element)
                rep = map(lambda element: 'IP' + element, rep)
                new_ip = ';'.join(rep)
                if r[14].value != new_ip:
                    self.textEdit.append(str(i + 3) + ' 行: ' + str(r[14].value) + ' 替换为 ' + new_ip)
                    r[14].value = new_ip
            i = i + 1
        try:
            wb.save(self.file_prj)
            self.textEdit.append('检查和更新完成.')
        except PermissionError:
            self.textEdit.append('写文件失败，关闭其他占用该文件的程序.' + self.file_prj)
        wb.close()

        xlApp = Dispatch("Excel.Application")
        xlApp.Visible = False
        xlApp.DisplayAlerts = False
        xlBook = xlApp.Workbooks.Open(self.file_prj)
        xlBook.Save()
        xlBook.Close()

    def checkpat2(self, doc_name, pat_name, pat_num):
        try:
            doc = Document(doc_name)
            found = False
            for i, para in enumerate(doc.tables[2].rows[4].cells[0].paragraphs):
                if pat_name in para.text:
                    found = True
                    result = re.search(pat_name + '.*号：' + pat_num, para.text)
                    if result is None:
                        self.textEdit.append(doc_name + ' 专利名和编号不匹配：' + pat_name + ' , ' + pat_num)
                        self.textEdit.append('文档内容：' + para.text)
                    break
            if not found:
                self.textEdit.append(doc_name + ' 全文找不到：' + pat_name)

        except PackageNotFoundError:
            self.textEdit.append('Error打开文件错误：' + doc_name)

    def upate_global(self):
        self.pat_dict.clear()
        self.pat_dict2.clear()
        wb = load_workbook(self.file_pat, read_only=True, data_only=True)
        ws = wb.active
        max_row_num = ws.max_row
        rangeCell = ws[f'A3:D{max_row_num}']
        for r in rangeCell:
            if r[0].value is None:
                break
            p_order = str(r[0].value).strip().zfill(2)
            p_name = str(r[1].value).strip()
            p_patnum = str(r[3].value).strip()
            self.pat_dict[p_name] = p_order
            self.pat_dict2[p_name] = p_patnum
        wb.close()

        self.arr_prj.clear()
        wb = load_workbook(self.file_prj, read_only=True, data_only=True)
        ws = wb.active
        if str(ws['A1'].value).find(u'公司') != -1:
            com_name = str(ws['A1'].value).split("公司")[0] + '公司'
        else:
            print("Error: 找不到 公司名")
        max_row_num = ws.max_row
        rangeCell = ws[f'A3:P{max_row_num}']
        for r in rangeCell:
            if r[0].value is None:
                break
            project = main.Project()
            project.p_comname = com_name
            project.p_order = str(r[0].value).strip().zfill(2)
            project.p_name = str(r[1].value).strip()
            project.p_start = r[2].value.strftime('%Y-%m-%d')
            project.p_end = r[3].value.strftime('%Y-%m-%d')
            project.p_cost = str(r[5].value).strip()
            project.p_people = str(r[6].value).strip()  # 人数
            project.p_owner = str(r[7].value).strip()  # 项目负责人
            project.p_rnd = str(r[8].value).strip()  # 研发人员
            project.p_money = str(r[9].value).strip()  # 总预算
            self.arr_prj.append(project)
        wb.close()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
